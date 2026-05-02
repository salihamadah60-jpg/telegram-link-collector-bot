"""
document_manager.py — Creates and manages the .docx files.

Two files are maintained:
  telegramLinks.docx  — collected Telegram group/channel invite links
  whatsappLinks.docx  — collected WhatsApp group invite links

Each file is regenerated fresh from the unsent links in the database
right before sending, so it always contains an accurate, deduplicated list.
"""

import logging
import os
from datetime import datetime, timezone
from typing import Literal

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import docx.opc.constants

import config
import database

logger = logging.getLogger("document_manager")

_TELEGRAM_FILE = "telegramLinks.docx"
_WHATSAPP_FILE = "whatsappLinks.docx"

_FILE_MAP: dict[str, str] = {
    "telegram": _TELEGRAM_FILE,
    "whatsapp": _WHATSAPP_FILE,
}


def _get_filename(link_type: Literal["telegram", "whatsapp"]) -> str:
    return _FILE_MAP[link_type]


def _build_document(
    links: list[dict],
    link_type: Literal["telegram", "whatsapp"],
) -> Document:
    """Build a styled Word document from a list of link dicts."""
    doc = Document()

    # Title
    title_text = (
        "Telegram Group & Channel Links"
        if link_type == "telegram"
        else "WhatsApp Group Invite Links"
    )
    title = doc.add_heading(title_text, level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}").italic = True
    meta.add_run(f"  |  Total links: {len(links)}").bold = True

    doc.add_paragraph()  # spacer

    # Links table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Link"
    hdr[2].text = "Date Collected"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for idx, link in enumerate(links, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)

        # Clickable hyperlink in cell [1]
        _add_hyperlink(row[1].paragraphs[0], link["url"], link["url"])

        date_str = ""
        if link.get("message_date"):
            d = link["message_date"]
            if hasattr(d, "strftime"):
                date_str = d.strftime("%Y-%m-%d")
        row[2].text = date_str

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run("End of file — collected by Telegram Userbot").italic = True

    return doc


def _add_hyperlink(paragraph, url: str, text: str):
    """Insert a clickable hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = docx.oxml.OxmlElement("w:r")
    rPr = docx.oxml.OxmlElement("w:rPr")

    rStyle = docx.oxml.OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    t = docx.oxml.OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


async def build_and_save(link_type: Literal["telegram", "whatsapp"]) -> str:
    """
    Fetch all unsent links from DB, build a .docx, save to disk.
    Returns the file path.
    """
    links = await database.get_unsent_links(link_type)
    if not links:
        raise ValueError(f"No unsent {link_type} links to write")

    doc = _build_document(links, link_type)
    filename = _get_filename(link_type)
    doc.save(filename)
    logger.info("Saved %d %s links to %s", len(links), link_type, filename)
    return filename


async def cleanup_file(link_type: Literal["telegram", "whatsapp"]) -> None:
    """Delete the .docx file from disk after sending."""
    filename = _get_filename(link_type)
    if os.path.exists(filename):
        os.remove(filename)
        logger.debug("Deleted %s from disk", filename)
